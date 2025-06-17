import docx
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False  # 正常显示负号
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import matplotlib
matplotlib.use('Agg')  # 使用非交互式绘图后端
if len(sys.argv) < 2:
    print("请提供Excel文件路径，例如： aa.exe data.xlsx")
    sys.exit(1)

excel_path = sys.argv[1]
if not os.path.exists(excel_path):
    print("文件不存在，请检查路径！")
    sys.exit(1)

# 使用 excel_path 读取 Excel 数据
df = pd.read_excel(excel_path)
df.columns = df.columns.str.strip()  # 清洗列名

# 初始化 Word 文档
doc = Document()
# 设置全局字体为仿宋
from docx.shared import Pt
from docx.oxml.ns import qn

style = doc.styles['Normal']
font = style.font
font.name = '仿宋_GB2312'
font.size = Pt(12)
font._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
title_para = doc.add_paragraph("全市电诈警情研判报告")
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.style = 'Title'  
doc.add_paragraph("（2025年1月1日至4月12日）")

# 一、总体情况
total_cases = len(df)
total_loss = df["涉案资金总和"].sum()
doc.add_heading("一、全市电诈警情总体情况", level=2)
doc.add_paragraph(f"2025年1月1日至4月12日，全市电诈警情{total_cases}起，损失金额{round(total_loss / 100000000, 2)}亿元。")

# 二、诈骗方式分析
doc.add_heading("二、从诈骗类别分析", level=2)
type_counts = df["诈骗方式"].value_counts()
type_percents = round(type_counts / total_cases * 100, 1)

doc.add_paragraph("（一）从各类别数量来看")
for scam_type, count in type_counts.items():
    percent = type_percents[scam_type]
    doc.add_paragraph(f"{scam_type} {count} 起，占比 {percent}%；")


type_counts = df['诈骗方式'].value_counts()
labels = type_counts.index.tolist()
sizes = type_counts.values.tolist()

fig, ax = plt.subplots(figsize=(6, 6))
wedges, texts, autotexts = ax.pie(
    sizes, autopct='%1.1f%%', startangle=140,
    textprops={'fontsize': 8}
)

ax.legend(wedges, labels, title="诈骗类型", loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
ax.set_title("诈骗类型案件数量占比", fontsize=12)
plt.tight_layout()
plt.savefig("诈骗类型_案件数量_饼图.png")
doc.add_picture("诈骗类型_案件数量_饼图.png", width=Inches(5.5))

doc.add_paragraph("（二）从各类别损失金额来看")
type_loss = df.groupby("诈骗方式")["涉案资金总和"].sum().sort_values(ascending=False)
type_loss_percent = round(type_loss / total_loss * 100, 1)
for scam_type, loss in type_loss.items():
    percent = type_loss_percent[scam_type]
    doc.add_paragraph(f"{scam_type} {round(loss / 10000, 2)} 万元，占比 {percent}%；")

loss_by_type = df.groupby("诈骗方式")["涉案资金总和"].sum().sort_values(ascending=False)
labels = loss_by_type.index.tolist()
sizes = loss_by_type.values.tolist()

fig, ax = plt.subplots(figsize=(6, 6))
wedges, texts, autotexts = ax.pie(
    sizes, autopct='%1.1f%%', startangle=140,
    textprops={'fontsize': 8}
)
ax.legend(wedges, labels, title="诈骗类型", loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
ax.set_title("诈骗类型涉案金额占比", fontsize=12)
plt.tight_layout()
plt.savefig("诈骗类型_涉案金额_饼图.png")
doc.add_picture("诈骗类型_涉案金额_饼图.png", width=Inches(5.5))

# 三、受骗群体分析
doc.add_heading("三、受骗群体分析", level=2)

doc.add_paragraph("（一）从性别看")
gender_counts = df["性别"].value_counts()
gender_percents = round(gender_counts / gender_counts.sum() * 100, 1)
for gender, count in gender_counts.items():
    doc.add_paragraph(f"{gender} {count} 起，占比 {gender_percents[gender]}%")

doc.add_paragraph("（一）从年龄结构看")
# 年龄段划分
bins = [0, 18, 35, 55, 150]
labels = ['18岁以下', '19至35岁', '36至55岁', '56岁以上']
df["年龄段"] = pd.cut(df["年龄"], bins=bins, labels=labels, right=True)
age_counts = df["年龄段"].value_counts().sort_index()
age_percents = round(age_counts / total_cases * 100, 1)

for age_group, count in age_counts.items():
    percent = age_percents[age_group]
    # 筛选该年龄段数据
    group_df = df[df["年龄段"] == age_group]
    type_counts = group_df["诈骗方式"].value_counts()
    type_percents = round(type_counts / type_counts.sum() * 100, 1)
    # 输出前 3 个诈骗方式
    top_types = [
        f"{t}{type_percents[t]}%" for t in type_counts.head(3).index
    ]
    doc.add_paragraph(f"{age_group}{count}起，占比{percent}%，高发手段：" + "、".join(top_types) + "；")

# 年龄结构饼状图
plt.figure(figsize=(6, 6))
plt.pie(age_counts, 
        labels=[f"{k} {v}起" for k, v in age_counts.items()], 
        autopct='%1.1f%%', 
        startangle=140, 
        textprops={'fontsize': 8})
plt.title("不同年龄段电诈案件占比")
plt.axis('equal')
plt.tight_layout()

age_pie_path = "年龄段_电诈占比_饼图.png"
plt.savefig(age_pie_path)

# 插入图表到 Word
doc.add_picture(age_pie_path, width=Inches(5.5))

doc.add_paragraph("（三）从职业类别看")
occupation_counts = df["身份"].fillna("其他").value_counts()
occupation_percents = round(occupation_counts / total_cases * 100, 1)

for occ, count in occupation_counts.head(10).items():
    percent = occupation_percents[occ]

    # 筛选该职业群体数据
    occ_df = df[df["身份"].fillna("其他") == occ]
    type_counts = occ_df["诈骗方式"].value_counts()
    if not type_counts.empty:
        top_type = type_counts.idxmax()
        top_percent = round(type_counts.max() / type_counts.sum() * 100, 1)
        doc.add_paragraph(f"{occ}{count}起，占比{percent}%，高发手段：{top_type}{top_percent}%；")

# 职业分布柱状图（前10职业）
top_occ = occupation_counts.head(10)
occ_names = top_occ.index.tolist()
occ_values = top_occ.values

plt.figure(figsize=(8, 5))
bars = plt.barh(occ_names, occ_values, color='mediumseagreen')
plt.xlabel("案件数量")
plt.title("职业群体电诈案件数量（前10）")
plt.gca().invert_yaxis()  # 最大的在最上

# 添加标签
for bar in bars:
    plt.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
             f'{bar.get_width()}', va='center', fontsize=8)

plt.tight_layout()
occupation_chart_path = "职业分布_柱状图.png"
plt.savefig(occupation_chart_path)

# 插入 Word 文档
doc.add_picture(occupation_chart_path, width=Inches(5.5))

# 四、发案区域分析
doc.add_heading("四、发案区域分析", level=2)
doc.add_paragraph("（一）各分、县（市）局警情情况")
district_counts = df["案发地区划"].value_counts()
district_percents = round(district_counts / total_cases * 100, 1)

for district, count in district_counts.items():
    percent = district_percents[district]
    sub_df = df[df["案发地区划"] == district]

    # 高发诈骗手段及占比
    scam_type_counts = sub_df["诈骗方式"].value_counts()
    scam_type_percents = round(scam_type_counts / len(sub_df) * 100, 1)
    top_scam_types = [
        f"{stype}{scam_type_percents[stype]}%" for stype in scam_type_counts.head(3).index
    ]

    # 高发人群及占比
    victim_counts = sub_df["身份"].fillna("其他").value_counts()
    victim_percents = round(victim_counts / len(sub_df) * 100, 1)
    top_victims = [
        f"{vtype}{victim_percents[vtype]}%" for vtype in victim_counts.head(3).index
    ]

    # 高发派出所及占比
    station_counts = sub_df["所属派出所"].fillna("其他").value_counts()
    station_percents = round(station_counts / len(sub_df) * 100, 1)
    top_stations = [
        f"{sname}{station_percents[sname]}%" for sname in station_counts.head(3).index
    ]

    # 输出汇总段落
    summary = f"{district}{count}起，占比{percent}%，" \
              f"高发手段（前三）：{', '.join(top_scam_types)}，" \
              f"高发人群：{', '.join(top_victims)}，" \
              f"高发派出所：{', '.join(top_stations)}；"
    doc.add_paragraph(summary)

district_counts = df["案发地区划"].value_counts()
labels = district_counts.index.tolist()
sizes = district_counts.values.tolist()

fig, ax = plt.subplots(figsize=(6, 6))
wedges, texts, autotexts = ax.pie(
    sizes, autopct='%1.1f%%', startangle=140,
    textprops={'fontsize': 8}
)
ax.legend(wedges, labels, title="地区", loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
ax.set_title("各地区电诈案件占比", fontsize=12)
plt.tight_layout()
plt.savefig("各地区_案件占比_饼图.png")
doc.add_picture("各地区_案件占比_饼图.png", width=Inches(5.5))

doc.add_paragraph("（二）全市电诈警情高发派出所")
top_stations = df["所属派出所"].fillna("其他").value_counts().head(10)
top_station_percents = round(top_stations / total_cases * 100, 1)

for idx, (station, count) in enumerate(top_stations.items(), start=1):
    percent = top_station_percents[station]
    sub_df = df[df["所属派出所"] == station]

    # 高发诈骗方式
    type_counts = sub_df["诈骗方式"].value_counts()
    type_percents = round(type_counts / len(sub_df) * 100, 1)
    top_types = [f"{t}{type_percents[t]}%" for t in type_counts.head(3).index]

    # 高发人群（所在单位）
    occ_counts = sub_df["身份"].fillna("其他").value_counts()
    occ_percents = round(occ_counts / len(sub_df) * 100, 1)
    top_occs = [f"{o}{occ_percents[o]}%" for o in occ_counts.head(3).index]

    summary = f"{idx}、{station}{count}起，占比{percent}%，" \
              f"高发手段：{', '.join(top_types)}，" \
              f"高发人群：{', '.join(top_occs)}；"
    doc.add_paragraph(summary)

# 生成前10派出所柱状图
top_stations = df["所属派出所"].fillna("其他").value_counts().head(10)
station_labels = top_stations.index.tolist()
station_counts = top_stations.values

plt.figure(figsize=(8, 5))
bars = plt.barh(station_labels, station_counts, color='skyblue')
plt.xlabel("案件数量")
plt.title("高发派出所警情数量（前10）")
plt.gca().invert_yaxis()  # 最大的排最上
for bar in bars:
    plt.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
             f'{bar.get_width()}', va='center', fontsize=8)
plt.tight_layout()
station_chart_path = "高发派出所_柱状图.png"
plt.savefig(station_chart_path)

# 插入图到 Word
doc.add_picture(station_chart_path, width=Inches(5.5))

# （三）各高校警情情况
doc.add_paragraph("（三）各高校警情情况")

# 过滤出“身份”为“学生”的记录
student_df = df[df["身份"] == "学生"]
student_count = len(student_df)
student_percent = round(student_count / total_cases * 100, 1)

doc.add_paragraph(f"学生警情共{student_count}起，占比{student_percent}%。")

# 统计学生群体高发诈骗方式
student_type_counts = student_df["诈骗方式"].value_counts()
student_type_percents = round(student_type_counts / student_count * 100, 1)

# 构造前 4 项文字
top_types = student_type_counts.head(4).index
top_strings = [
    f"{stype}{student_type_percents[stype]}%" for stype in top_types
]
doc.add_paragraph("高发手段：" + "、".join(top_strings) + "。")

# 学生群体高发诈骗方式饼状图
top_n = 6  # 显示前6种，其余归为“其他”
student_type_counts = student_df["诈骗方式"].value_counts()
student_top_types = student_type_counts.head(top_n)
other_count = student_type_counts[top_n:].sum()

labels = list(student_top_types.index) + (["其他"] if other_count > 0 else [])
sizes = list(student_top_types.values) + ([other_count] if other_count > 0 else [])

fig, ax = plt.subplots(figsize=(6, 6))
wedges, texts, autotexts = ax.pie(
    sizes, autopct='%1.1f%%', startangle=140,
    textprops={'fontsize': 8}
)
ax.legend(wedges, labels, title="诈骗方式", loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
ax.set_title("学生群体诈骗方式占比", fontsize=12)
plt.tight_layout()
plt.savefig("学生_诈骗方式_饼图.png")
doc.add_picture("学生_诈骗方式_饼图.png", width=Inches(5.5))

# 提取高校记录：所在单位包含“大学”或“学院”
university_df = student_df[student_df["所在单位"].fillna("").str.contains("大学|学院")]
university_counts = university_df["所在单位"].value_counts()
university_total = university_counts.sum()

doc.add_paragraph("高发学校：")
for school, count in university_counts.head(10).items():
    percent = round(count / student_count * 100, 1)
    doc.add_paragraph(f"{school}{count}起，占比{percent}%，")

# 生成高校柱状图
top_univ = university_counts.head(10)
schools = top_univ.index.tolist()
counts = top_univ.values

plt.figure(figsize=(8, 5))
bars = plt.barh(schools, counts, color='lightgreen')
plt.xlabel("案件数量")
plt.title("高发高校警情数量（前10）")
plt.gca().invert_yaxis()  # 最大的在最上

# 添加文本标签
for bar in bars:
    plt.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
             f'{bar.get_width()}', va='center', fontsize=8)

plt.tight_layout()
school_chart_path = "高发高校_柱状图.png"
plt.savefig(school_chart_path)

# 插入 Word 文档
doc.add_picture(school_chart_path, width=Inches(5.5))

# 高发区域统计
doc.add_paragraph("高发区域：")
area_counts = df["案发地区划"].value_counts().head(5)
area_percents = round(area_counts / total_cases * 100, 1)
for area, count in area_counts.items():
    percent = area_percents[area]
    doc.add_paragraph(f"{area}{count}起、占比{percent}%，")

# 高发派出所统计
doc.add_paragraph("高发派出所：")
station_counts = df["所属派出所"].fillna("其他").value_counts().head(5)
station_percents = round(station_counts / total_cases * 100, 1)
for station, count in station_counts.items():
    percent = station_percents[station]
    doc.add_paragraph(f"{station}{count}起、占比{percent}%，")

# 学生警情高发派出所柱状图
student_station_counts = student_df["所属派出所"].fillna("其他").value_counts().head(10)
stations = student_station_counts.index.tolist()
counts = student_station_counts.values

plt.figure(figsize=(8, 5))
bars = plt.barh(stations, counts, color='orange')
plt.xlabel("案件数量")
plt.title("学生警情高发派出所（前10）")
plt.gca().invert_yaxis()  # 最大的排最上
for bar in bars:
    plt.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
             f'{bar.get_width()}', va='center', fontsize=8)

plt.tight_layout()
student_station_chart = "学生警情_高发派出所.png"
plt.savefig(student_station_chart)

# 插入图表到 Word
doc.add_picture(student_station_chart, width=Inches(5.5))    

# 自动识别所有诈骗途径相关字段
channel_columns = [col for col in df.columns if '诈骗途径' in col]

# 引流方式分析部分
doc.add_heading("五、引流途径分析", level=2)
doc.add_paragraph("（一）从诈骗途径来看")

for col in channel_columns:
    way = col.replace('诈骗途径：', '')
    count = df[col].fillna('').apply(lambda x: str(x).strip() != '').sum()
    percent = round(count / total_cases * 100, 1)
    doc.add_paragraph(f"通过{way}引流诈骗 {count} 起，占比 {percent}%")

# 保存报告
doc.save("全市电诈警情研判报告_完整版.docx")
print("✅ 生成完成：全市电诈警情研判报告_完整版.docx")